"""
DOCX loader: extract text and tables from native .docx files using python-docx.
Produces `layout_data` compatible with `flatten_layout` and `table_data` as list of rows.
"""
from io import BytesIO
from typing import List

try:
    from docx import Document
except Exception:
    Document = None

def extract_docx(docx_bytes: bytes):
    if Document is None:
        raise RuntimeError("python-docx is not installed. Install with `pip install python-docx`")

    doc = Document(BytesIO(docx_bytes))

    # Build a single-page layout_data structure compatible with flatten_layout
    # layout_data -> list[page] -> list[block] -> { 'lines': [ { 'spans': [ {'text': ...} ] } ] }
    page_blocks = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        block = {
            "lines": [
                {"spans": [{"text": text}]}
            ]
        }
        page_blocks.append(block)

    # Extract tables
    table_rows: List[List[str]] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append(cells)

    layout_data = [page_blocks]
    return layout_data, table_rows
